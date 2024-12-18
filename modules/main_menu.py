#Импорт кастом тк интера
import customtkinter as ctk
from modules.button_creation import CustomButton
from modules.label_creation import CustomLabel
from modules.new_windows import ToplevelWindow
from PIL import Image, ImageTk
from tkinter import filedialog
from docx import Document
import PyPDF2
from modules.ai import client
from modules.plagiat_checker import detector
from modules.lecs_checker import assess_text_originality
from difflib import SequenceMatcher
from modules.test_plagiat import wiki_wiki

# Создание констант с размерами окна
APP_WIDTH = 1280
APP_HEIGHT = 832
# Клас настройки окна 
class App(ctk.CTk):
    # метод конструктор для класса
    def __init__(self, fg_color):
        # Получаем свойство цвета заднего фона
        super().__init__(fg_color)
        # Свойства для размера окна и его расположения
        self.APP_WIDTH = APP_WIDTH
        self.APP_HEIGHT = APP_HEIGHT
        self.X = 200
        self.Y = 0
        
        # Задаем размеры окна и его расположение
        self.geometry(f"{self.APP_WIDTH}x{self.APP_HEIGHT}+{self.X}+{self.Y}")
        # Говорим что нельзя менять размер окна
        self.resizable(False,False)
        # Указываем название окна
        self.title("makaki")
        
        self.main_frame = ctk.CTkFrame(self, fg_color="#7593EB")
        self.main_frame.pack(fill="both", expand=True)
         # Загрузка изображения для кнопки
         
        
        self.button1 = CustomButton(self.main_frame, command = self.new_window_with_text, text="TEXT")
        self.button1.place(x= 530, y=650)
        
        
        self.button2 = CustomButton(self.main_frame, text="FILE", command = self.new_window_with_file)
        self.button2.place(x= 1010, y=650)

         
        self.button3 = CustomButton(self.main_frame, text="IMAGE", command = self.new_window_with_image)
        self.button3.place(x= 890, y=720)

        self.button_image = ctk.CTkImage(Image.open("modules\\images\\image.png"), size=(200,30))
        self.button4 = CustomButton(self.main_frame, text="COMPARE", command=self.new_window_with_comparison)
        self.button4.place(x= 640, y=720)

        self.bg_image = Image.open("modules\\images\\new_menu_frame.png")
        self.bg_image= self.bg_image.resize((self.winfo_width(), self.winfo_height()))
        self.bg_photo = ImageTk.PhotoImage(self.bg_image)
        
        # self.label1 = CustomLabel(self, image = self.bg_photo, text="")
        # self.label1.place(relx=0, rely=0, relwidth = 1, relheight =1)
        
        canvas = ctk.CTkCanvas(self.main_frame, width=600, height=400, highlightthickness=0)
        canvas.pack(fill="both", expand=True)
        canvas.create_image(0, 0, image=self.bg_photo, anchor="nw")
        
        # self.attributes("-transparentcolor", "white")
        
        # Полнимаем кнопки
        self.button1.lift()
        self.button2.lift()
        self.button3.lift()
        self.button4.lift()

    def new_window_with_comparison(self):
            self.toplevel_window_with_comparison = ToplevelWindow(self) 
            self.toplevel_window_with_comparison.geometry(f"{self.APP_WIDTH}x{self.APP_HEIGHT}+{self.X}+{self.Y}")
            self.withdraw()

            self.image_back= ctk.CTkImage(Image.open("modules\\images\\back_for_comparison.png"), size= (1280, 832))
            self.label_image4= ctk.CTkLabel(self.toplevel_window_with_comparison, image= self.image_back, text="")
            self.label_image4.place(x= 0, y=0)

            self.button_image_start= ctk.CTkImage(Image.open("modules\\images\\button_start.png"), size=(282, 64))
            self.button_start4 = CustomButton(self.toplevel_window_with_comparison,fg_color="#282C43" ,image= self.button_image_start, command=self.new_window_with_comparison_result)
            self.button_start4.place(x = 489, y=706)
            self.button_start4.lift()

            self.input_text4_1 = ctk.CTkTextbox(self.toplevel_window_with_comparison, width=451, height=516, fg_color="#D4DEE6", text_color="black", wrap="word")
            self.input_text4_1.place(x = 80, y=140)
            self.input_text4_1.lift()

            self.input_text4_2 = ctk.CTkTextbox(self.toplevel_window_with_comparison, width=451, height=516, fg_color="#D4DEE6", text_color="black", wrap="word")
            self.input_text4_2.place(x = 750, y=140)
            self.input_text4_2.lift()

            self.home_image= ctk.CTkImage(Image.open("modules\\images\\home.png"), size=(100, 20))
            self.comeback4_1_button=CustomButton(self.toplevel_window_with_comparison,width=100, height=20, image = self.home_image, fg_color="#D4DEE6", command = self.comeback4_1 )
            self.comeback4_1_button.place(x=50, y=70)
            self.comeback4_1_button.lift()

            self.download_image= ctk.CTkImage(Image.open("modules\\images\\download.png"), size=(50, 50))
            self.upload4_1_button=CustomButton(self.toplevel_window_with_comparison,width=50, height=50,image=self.download_image, fg_color="#D4DEE6", command=self.load_file4_1 )
            self.upload4_1_button.place(x=80, y=720)
            self.upload4_1_button.lift()

            self.upload4_2_button=CustomButton(self.toplevel_window_with_comparison,width=50, height=50,image=self.download_image, fg_color="#D4DEE6", command=self.load_file4_2 )
            self.upload4_2_button.place(x=1137, y=720)
            self.upload4_2_button.lift()
            
    def comeback4_1(self):
            self.toplevel_window_with_comparison.destroy()
            self.deiconify()  

    def load_file4_1(self):
            file_path = filedialog.askopenfilename(title="Выберите файл",filetypes=[("Text Files", "*.txt"),("PDF Files", "*.pdf"), ("Word Documents", "*.docx"), ("All Files", "*.*")])
            if file_path:
                try:
                    if file_path.endswith(".txt"):
                        # Чтение текстового файла
                        with open(file_path, "r", encoding="utf-8") as file:
                            content = file.read()
                    elif file_path.endswith(".docx"):
                        # Чтение Word-документа
                        doc = Document(file_path)
                        content = "\n".join([para.text for para in doc.paragraphs])
                    elif file_path.endswith(".pdf"):
                
                        with open(file_path, "rb") as file:
                            reader = PyPDF2.PdfReader(file)
                            content = ""
                            for page in reader.pages:
                                content += page.extract_text()
                    else:
                        print("Неподдерживаемый формат файла")
                        return

                    # Очистка и вставка текста в input_text3
                    self.input_text4_1.delete("1.0", "end")  # Очистить текстовое поле
                    self.input_text4_1.insert("1.0", content)  # Вставить текст

                except Exception as e:
                    print(f"Ошибка при чтении файла: {e}")

    def load_file4_2(self):
            file_path = filedialog.askopenfilename(title="Выберите файл",filetypes=[("Text Files", "*.txt"),("PDF Files", "*.pdf"), ("Word Documents", "*.docx"), ("All Files", "*.*")])
            if file_path:
                try:
                    if file_path.endswith(".txt"):
                        # Чтение текстового файла
                        with open(file_path, "r", encoding="utf-8") as file:
                            content = file.read()
                    elif file_path.endswith(".docx"):
                        # Чтение Word-документа
                        doc = Document(file_path)
                        content = "\n".join([para.text for para in doc.paragraphs])
                    elif file_path.endswith(".pdf"):
                
                        with open(file_path, "rb") as file:
                            reader = PyPDF2.PdfReader(file)
                            content = ""
                            for page in reader.pages:
                                content += page.extract_text()
                    else:
                        print("Неподдерживаемый формат файла")
                        return

                    # Очистка и вставка текста в input_text3
                    self.input_text4_2.delete("1.0", "end")  # Очистить текстовое поле
                    self.input_text4_2.insert("1.0", content)  # Вставить текст

                except Exception as e:
                    print(f"Ошибка при чтении файла: {e}")

    def new_window_with_comparison_result(self):
            self.toplevel_window_with_comparison_result = ToplevelWindow(self) 
            self.toplevel_window_with_comparison_result.geometry(f"{self.APP_WIDTH}x{self.APP_HEIGHT}+{self.X}+{self.Y}")    
            if not self.input_text4_1.get("1.0","end").strip() or not self.input_text4_2.get("1.0","end").strip():
                self.toplevel_window_with_comparison_result.destroy()
                return
            text4_1 = self.input_text4_1.get("1.0","end")
            text4_2 = self.input_text4_2.get("1.0","end")  
            match = SequenceMatcher(None, text4_1, text4_2)

            result = match.ratio() * 100           
            self.toplevel_window_with_comparison.withdraw()

            self.image_back_comparison_result= ctk.CTkImage(Image.open("modules\\images\\comparison_result.png"), size= (1280, 832))
            self.label_image_comparison_result= ctk.CTkLabel(self.toplevel_window_with_comparison_result, image= self.image_back_comparison_result, text="")
            self.label_image_comparison_result.place(x= 0, y=0)

            self.input_text4_1_result = ctk.CTkTextbox(self.toplevel_window_with_comparison_result, width=451, height=516, fg_color="#D4DEE6", text_color="black", wrap="word")
            self.input_text4_1_result.place(x = 80, y=140)
            self.input_text4_1_result.insert("1.0", text4_1)
            self.input_text4_1_result.lift()

            self.input_text4_2_result = ctk.CTkTextbox(self.toplevel_window_with_comparison_result, width=451, height=516, fg_color="#D4DEE6", text_color="black", wrap="word")
            self.input_text4_2_result.place(x = 750, y=140)
            self.input_text4_2_result.insert("1.0", text4_2)
            self.input_text4_2_result.lift()

            self.input_text4_3_result = ctk.CTkTextbox(self.toplevel_window_with_comparison_result, width=282, height=64, fg_color="#D4DEE6", text_color="black", wrap="word")
            self.input_text4_3_result.place(x = 489, y=706)
            self.input_text4_3_result.insert("1.0", f'Процент плагиата: {result:.2f}%')
            self.input_text4_3_result.lift()

            self.home_image= ctk.CTkImage(Image.open("modules\\images\\home.png"), size=(100, 20))
            self.comeback4_1_button_result=CustomButton(self.toplevel_window_with_comparison_result,width=100, height=20, image = self.home_image, fg_color="#D4DEE6", command = self.comeback4_1_result )
            self.comeback4_1_button_result.place(x=50, y=70)
            self.comeback4_1_button_result.lift()

    def comeback4_1_result(self):
            self.toplevel_window_with_comparison_result.destroy()
            self.deiconify()     

        

    def new_window_with_text(self):
            
            self.toplevel_window_with_text = ToplevelWindow(self) 
            self.toplevel_window_with_text.geometry(f"{self.APP_WIDTH}x{self.APP_HEIGHT}+{self.X}+{self.Y}")
            self.withdraw()
            
            self.button_image_start= ctk.CTkImage(Image.open("modules\\images\\button_start.png"), size=(282, 64))
            self.button_start = CustomButton(self.toplevel_window_with_text, fg_color="#282C43" ,image= self.button_image_start, command=self.new_window_with_text_result)
            self.button_start.place(x = 499, y=666)

            self.home_image= ctk.CTkImage(Image.open("modules\\images\\home.png"), size=(100, 20))
            self.comeback3_button=CustomButton(self.toplevel_window_with_text,width=100, height=20, image = self.home_image, fg_color="#D4DEE6", command = self.comeback3 )
            self.comeback3_button.place(x=50, y=70)

            self.image_back= ctk.CTkImage(Image.open("modules\\images\\background_image_for_second_frame.png"), size= (1280, 832))
            self.label_image= ctk.CTkLabel(self.toplevel_window_with_text, image= self.image_back, text="")
            self.label_image.place(x= 0, y=0)

            self.input_text_window= ctk.CTkImage(Image.open("modules\\images\\text_input.png"), size=(451, 516))
            self.input_text = ctk.CTkTextbox(self.toplevel_window_with_text, width=451, height=516, fg_color="#D4DEE6", text_color="black", wrap="word")
            self.input_text.place(x = 414, y=85)
            self.input_text.bind("<Control-v>", self.paste_text)
            self.input_text.bind("<Key>", self.limit_text)
            
            self.comeback3_button.lift()
            self.input_text.focus_set()
            self.button_start.lift()
            self.input_text.lift()
    def comeback3(self):
            self.toplevel_window_with_text.destroy()
            self.deiconify()

    def limit_text(self, event):
        # Дозволяємо тільки введення до 512 символів
        current_text = self.input_text.get("1.0", "end").strip()
        if len(current_text) >= 512 and event.keysym not in ("BackSpace", "Delete", "Left", "Right", "Up", "Down"):
            # Забороняємо введення додаткових символів
            return "break"
            
    def paste_text(self, event=None):
            # Отримуємо текст із буфера обміну
            clipboard_text = self.clipboard_get()
    
            # Поточний текст у віджеті
            current_text = self.input_text.get("1.0", "end").strip()
    
            # Перевіряємо, чи вставка перевищить обмеження
            if len(current_text) + len(clipboard_text) > 512:
                #Вираховуємо кількість дозволених символів для вставки
                allowed_length = 512 - len(current_text)
        
                #Вставляємо лише допустиму кількість символів
                self.input_text.insert("insert", clipboard_text[:allowed_length])
            else:
                # Вставляємо весь текст з буфера обміну
                self.input_text.insert("insert", clipboard_text)
            print (len(clipboard_text))
            return "break"
            
    def new_window_with_image(self):
        self.toplevel_window_with_image = ToplevelWindow(self, width= 451, height= 516) 
        self.toplevel_window_with_image.geometry(f"{self.APP_WIDTH}x{self.APP_HEIGHT}+{self.X}+{self.Y}")
        self.withdraw()
        
        self.button_image_start2= ctk.CTkImage(Image.open("modules\\images\\button_start.png"), size=(282, 64))
        self.button_start2 = CustomButton(self.toplevel_window_with_image, fg_color="#282C43", image= self.button_image_start2, command = self.new_window_with_image_result)
        self.button_start2.place(x = 499, y=666)
  
        self.input_text_window2= ctk.CTkImage(Image.open("modules\\images\\text_input.png"), size=(451, 516))
        self.input_text2 = ctk.CTkEntry(self.toplevel_window_with_image, width=451, height= 516, fg_color="#D4DEE6")
        self.input_text2.place(x = 415.5, y=80)
        self.input_text2.place(x = 414, y=30)
        self.input_text2.place(x = 415.5, y=85)
        
        self.image_back2= ctk.CTkImage(Image.open("modules\\images\\back_for_image.png"), size= (1280, 832))
        self.label_image= ctk.CTkLabel(self.toplevel_window_with_image, image= self.image_back2, text="")
        self.label_image.place(x= 0, y=0)
        
        self.home_image= ctk.CTkImage(Image.open("modules\\images\\home.png"), size=(100, 20))
        self.comeback2_button=CustomButton(self.toplevel_window_with_image,width=100, height=20, image = self.home_image, fg_color="#D4DEE6", command = self.comeback2 )
        self.comeback2_button.place(x=50, y=70)

        self.download_image= ctk.CTkImage(Image.open("modules\\images\\download.png"), size=(50, 50))
        self.upload_button2=CustomButton(self.toplevel_window_with_image, fg_color="#D4DEE6",image=self.download_image,width=50, height=50, command=self.load_image )
        self.upload_button2.place(x=840, y=680)
        
        self.comeback2_button.lift()
        self.button_start2.lift()
        self.input_text2.lift()
        self.upload_button2.lift()
    def comeback2(self):
        self.toplevel_window_with_image.destroy()
        self.deiconify()

    def load_image(self):
        image_path = filedialog.askopenfilename(
        title="Выберите фото",
        filetypes=[("Image", "*.png"), ("All Files", "*.*")])
        if image_path:
            try:
                image = Image.open(image_path)
                self.input_text_window2 = ctk.CTkImage(image, size=(451, 516))
            except Exception as e:
                print(f"Ошибка при чтении файла: {e}")
        self.input_text2 = ctk.CTkLabel(self.toplevel_window_with_image, image=self.input_text_window2, text="")
        self.input_text2.place(x=414.5, y=85)
        self.label1_image=ctk.CTkLabel(self.toplevel_window_with_image,)
        self.label1_image.place(x = 0, y=0)




    def new_window_with_file(self):

        self.toplevel_window_with_file = ToplevelWindow(self) 
        self.toplevel_window_with_file.geometry(f"{self.APP_WIDTH}x{self.APP_HEIGHT}+{self.X}+{self.Y}")
        self.withdraw()

        self.image_back3= ctk.CTkImage(Image.open("modules\\images\\back_for_file.png"), size= (1280, 832))
        self.label_image= ctk.CTkLabel(self.toplevel_window_with_file, image= self.image_back3, text="")
        self.label_image.place(x= 0, y=0)

        self.button_image_start3= ctk.CTkImage(Image.open("modules\\images\\button_start.png"), size=(282, 64))
        self.button_start3 = CustomButton(self.toplevel_window_with_file,fg_color="#282C43", image= self.button_image_start3, command=self.new_window_with_file_result )
        self.button_start3.place(x = 499, y=666)
        self.button_start3.lift()
        
        self.input_text3 = ctk.CTkTextbox(self.toplevel_window_with_file, width=451, height=516, fg_color="#D4DEE6", text_color="black", wrap="word")
        self.input_text3.place(x = 414, y=85)
        self.input_text3.lift()

        self.home_image= ctk.CTkImage(Image.open("modules\\images\\home.png"), size=(100, 20))
        self.comeback_button=CustomButton(self.toplevel_window_with_file,width=100, height=20, image = self.home_image, fg_color="#D4DEE6", command = self.comeback )
        self.comeback_button.place(x=50, y=70)
        self.comeback_button.lift()

        self.download_image= ctk.CTkImage(Image.open("modules\\images\\download.png"), size=(50, 50))
        self.upload_button=CustomButton(self.toplevel_window_with_file,width=50, height=50,image=self.download_image, fg_color="#D4DEE6", command=self.load_file )
        self.upload_button.place(x=840, y=680)
        self.upload_button.lift()

    def comeback(self):
            self.toplevel_window_with_file.destroy()
            self.deiconify()

    def load_file(self):
            file_path = filedialog.askopenfilename(title="Выберите файл",filetypes=[("Text Files", "*.txt"),("PDF Files", "*.pdf"), ("Word Documents", "*.docx"), ("All Files", "*.*")])
            if file_path:
                try:
                    if file_path.endswith(".txt"):
                        # Чтение текстового файла
                        with open(file_path, "r", encoding="utf-8") as file:
                            self.content = file.read()
                    elif file_path.endswith(".docx"):
                        # Чтение Word-документа
                        doc = Document(file_path)
                        self.content = "\n".join([para.text for para in doc.paragraphs])
                    elif file_path.endswith(".pdf"):
                
                        with open(file_path, "rb") as file:
                            reader = PyPDF2.PdfReader(file)
                            self.content = ""
                            for page in reader.pages:
                                self.content += page.extract_text()
                    else:
                        print("Неподдерживаемый формат файла")
                        return
                    self.content = self.content[:512]

                    # Очистка и вставка текста в input_text3
                    self.input_text3.delete("1.0", "end")  # Очистить текстовое поле
                    self.input_text3.insert("1.0", self.content)  # Вставить текст

                except Exception as e:
                    print(f"Ошибка при чтении файла: {e}")
    
    def new_window_with_text_result(self):
        text = self.input_text.get("1.0","end")  
        
        response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": f"Привет ты учитель и проверяешь учеников на честность , они написали текст сами или текст написала нейросеть . Тебе нужно написать, вероятность написания текста нейросетью , вероятность написания человеком ,вероятность что человек взял текст из интернета  .Также если видишь что текст похож на научный (слишком много цифр ) то скорее всего этот текст взят из интернета . Ответ напиши в процентах. Вот сам текст: {text}"}]
)
        self.result_ai = response.choices[0].message.content
        print(self.result_ai)
        
        result = detector(text)


        self.result_plagiat= result[0]['score']  
        self.result_plagiat = int(self.result_plagiat)*100
        
        wiki_theme = text.split(" ")[0]
        print(wiki_theme)
        page = wiki_wiki.page(wiki_theme)
        print(page.text[:500])
        match = SequenceMatcher(None, text, page.text[:500])

        result = match.ratio() * 100 
        
        print(result)


        self.final_result = f"Перевірено ШІ: {self.result_ai} \nПеревірено на ознаки автоматичної генерації: {self.result_plagiat}% \nПеревірено на лексичну різноманітність: {assess_text_originality(text)} \n 'Відсоток плагіату: {result:.2f}%'"
        
        self.toplevel_window_with_text_result = ToplevelWindow(self) 
        self.toplevel_window_with_text_result.geometry(f"{self.APP_WIDTH}x{self.APP_HEIGHT}+{self.X}+{self.Y}")    
          
                 
        self.toplevel_window_with_text.withdraw()

        self.image1_back_result= ctk.CTkImage(Image.open("modules\\images\\result_bg.jpg"), size= (1280, 832))
        self.label1_image_result=ctk.CTkLabel(self.toplevel_window_with_text_result, image= self.image1_back_result, text="" )
        self.label1_image_result.place(x = 0, y=0)

        self.input_text1_result = ctk.CTkTextbox(self.toplevel_window_with_text_result, width=451, height=516, fg_color="#D4DEE6", text_color="black", wrap="word")
        self.input_text1_result.place(x = 415.5, y=85)
        self.input_text1_result.insert("1.0", text)
        self.input_text1_result.lift()

        self.home_image= ctk.CTkImage(Image.open("modules\\images\\home.png"), size=(100, 20))
        self.comeback4_button=CustomButton(self.toplevel_window_with_text_result,width=100, height=20, image = self.home_image, fg_color="#D4DEE6", command = self.comeback4 )
        self.comeback4_button.place(x=50, y=70)
        self.comeback4_button.lift()

        self.input_text11_result = ctk.CTkTextbox(self.toplevel_window_with_text_result, width=451, height=100, fg_color="#D4DEE6", text_color="black", wrap="word")
        self.input_text11_result.place(x = 415.5, y=666)
        self.input_text11_result.insert("1.0", self.final_result)
    
        self.input_text11_result.lift()
    def comeback4 (self):
            self.toplevel_window_with_text_result.destroy()
            self.deiconify()

    def new_window_with_image_result(self):
        self.toplevel_window_with_image_result = ToplevelWindow(self) 
        self.toplevel_window_with_image_result.geometry(f"{self.APP_WIDTH}x{self.APP_HEIGHT}+{self.X}+{self.Y}")     
        
        self.toplevel_window_with_image.withdraw()

        self.image3_back_result= ctk.CTkImage(Image.open("modules\\images\\result_bg.jpg"), size= (1280, 832))
        self.label3_image_result=ctk.CTkLabel(self.toplevel_window_with_image_result, image= self.image3_back_result, text="" )
        self.label3_image_result.place(x = 0, y=0)

        self.input_text2_result = ctk.CTkLabel(self.toplevel_window_with_image_result, width=451, height= 516, fg_color="#D4DEE6",image=self.input_text_window2, text="")
        self.input_text2_result.place(x = 415.5, y=85)
        self.input_text2_result.lift()

        self.home_image= ctk.CTkImage(Image.open("modules\\images\\home.png"), size=(100, 20))
        self.comeback5_button=CustomButton(self.toplevel_window_with_image_result,width=100, height=20, image = self.home_image, fg_color="#D4DEE6", command = self.comeback5 )
        self.comeback5_button.place(x=50, y=70)
        self.comeback5_button.lift()

        self.input_text21_result = ctk.CTkEntry(self.toplevel_window_with_image_result, width=451, height= 100, fg_color="#D4DEE6")
        self.input_text21_result.place(x = 415.5, y=666)
        self.input_text21_result.lift()
    def comeback5 (self):
            self.toplevel_window_with_image_result.destroy()
            self.deiconify()


    def new_window_with_file_result(self):
        text_ = self.input_text3.get("1.0","end")   
        response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": f"Привет ты учитель и проверяешь учеников на честность , они написали текст сами или текст написала нейросеть . Ты сможешь проверить  текст ученика ?  Если тогда можешь давать ответ одним словом да или нет , вероятность написания текста нейросетью , вероятность написания человеком ,вероятность что человек взял текст из интернета  . Молодец , можно я тебя немного исправлю . Если ты видишь художественные произведения (сказки ,фантастика и подобные ) скорее всего написал человек .Также если видишь что текст похож на научный (слишком много цифр ) то скорее всего этот текст взят из интернета . Ответ давай в процентах. Этот текст написала нейросеть? Вот сам текст: {text_}"}]
)
        self.result_ai_file = response.choices[0].message.content
        print(self.result_ai_file)
        
        result_ = detector(text_)
        self.result_plagiat_file= result_[0]['score']  
        self.result_plagiat_file = float(self.result_plagiat_file)*100


        self.final_result_file = f"Перевірено ШІ: {self.result_ai_file} \nПеревірено на ознаки автоматичної генерації: {self.result_plagiat_file}% \nПеревірено на лексичну різноманітність: {assess_text_originality(text_)}"
        self.toplevel_window_with_result = ToplevelWindow(self) 
        self.toplevel_window_with_result.geometry(f"{self.APP_WIDTH}x{self.APP_HEIGHT}+{self.X}+{self.Y}") 
                     
        self.toplevel_window_with_file.withdraw()

        self.image2_back_result= ctk.CTkImage(Image.open("modules\\images\\result_bg.jpg"), size= (1280, 832))
        self.label2_image_result=ctk.CTkLabel(self.toplevel_window_with_result, image= self.image2_back_result, text="" )
        self.label2_image_result.place(x = 0, y=0)

        self.input_text3_result = ctk.CTkTextbox(self.toplevel_window_with_result, width=451, height=516, fg_color="#D4DEE6", text_color="black", wrap="word")
        self.input_text3_result.place(x = 415.5, y=85)
        self.input_text3_result.insert("1.0", text_)
        self.input_text3_result.lift()

        self.input_text33_result = ctk.CTkTextbox(self.toplevel_window_with_result, width=451, height=100, fg_color="#D4DEE6", text_color="black", wrap="word")
        self.input_text33_result.place(x = 415.5, y=666)
        self.input_text33_result.insert("1.0", self.final_result_file)
        self.input_text33_result.lift()

        self.home_image= ctk.CTkImage(Image.open("modules\\images\\home.png"), size=(100, 20))
        self.comeback6_button=CustomButton(self.toplevel_window_with_result,width=100, height=20, image = self.home_image, fg_color="#D4DEE6", command = self.comeback6 )
        self.comeback6_button.place(x=50, y=70)
        self.comeback6_button.lift()
    def comeback6(self):
            self.toplevel_window_with_result.destroy()
            self.deiconify()
        
        

   

       
       

# Создаем окно и задаем ему цвет 
app = App(fg_color = "white")
